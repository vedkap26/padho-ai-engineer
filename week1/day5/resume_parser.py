import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq

from pydantic import BaseModel, Field
load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("GROQ_API_KEY environment variable is not set.")
client = Groq(api_key=my_api_key)
model = "llama-3.3-70b-versatile"

job_description = """Description
At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. The focus we have on our customers is why we are one of the world’s most beloved brands – customer obsession is part of our company DNA. Our Software Development Engineers (SDEs) use latest technology to solve complex problems and get to see the impact of their work first-hand. The challenges SDEs solve for at Amazon are big and influence millions of customers, sellers, and products around the world. We are looking for individuals who are passionate about creating new products, features, and services from scratch while managing ambiguity and the pace of a company where development cycles are measured in weeks, not years. If this sounds interesting to you, apply and come chart your own path at Amazon.
Applications are reviewed on a rolling basis. For an update on your status, or to confirm your application was submitted successfully, please login to your candidate portal. NOTE: Amazon works with a high volume of applicants, so we appreciate your patience as we review applications.

Key job responsibilities
• Collaborate with experienced cross-disciplinary Amazonians to conceive, design, and bring innovative products and services to market.
• Design and build innovative technologies in a large distributed computing environment and help lead fundamental changes in the industry.
• Create solutions to run predictions on distributed systems with exposure to innovative technologies at incredible scale and speed.
• Build distributed storage, index, and query systems that are scalable, fault-tolerant, low cost, and easy to manage/use.
• Design and code the right solutions starting with broadly defined problems.
• Work in an agile environment to deliver high-quality software.

Note: This role is part of the rekindle program. For more details on rekindle program, please visit - https://www.amazon.jobs/en/landing_pages/rekindle

Basic Qualifications
- 1+ years of non-internship professional software development experience
- Experience programming with at least one software programming language

Preferred Qualifications
- Bachelor's degree in computer science or equivalent

Our inclusive culture empowers Amazonians to deliver the best results for our customers. If you have a disability and need a workplace accommodation or adjustment during the application and hiring process, including support for the interview or onboarding process, please visit https://amazon.jobs/content/en/how-we-hire/accommodations for more information. If the country/region you’re applying in isn’t listed, please contact your Recruiting Partner."""


# part 1 creating job schema
class JobDescription(BaseModel):
    role:str
    required_skills: list[str] 
    preferred_skills: list[str]
    min_experience: float|None = None
    education_requirement: str|None = None
    responsibilities: list[str]
job_schema = JobDescription.model_json_schema()

system_prompt=f"""please extract the job description information and give a json output matching this schema: {job_schema}. Only include fields you can actually find in the text; omit fields that aren't present. important: do not return the actual schema itself do not reuturn the field like "properties","title".\n\ntext: {job_description}"""
user_prompt=f"""analyze the job description and extract the required information. Return the output in JSON format matching the schema provided. Only include fields that are present in the text; omit any fields that aren't found. Do not return the actual schema itself or any field names like "properties" or "title".\n\ntext: {job_description}"""
message_system = {"role": "system", "content": system_prompt}
message_user = {"role": "user", "content": user_prompt}
response_format = {
    "type": "json_object"     
}
messages=[message_system, message_user]
response = client.chat.completions.create(model=model, messages=messages,response_format=response_format)
answer = response.choices[0].message.content
raw_data = answer

import json
job_file=json.loads(raw_data)

job_description = JobDescription(**job_file)

print(job_description.role)
print(job_description.required_skills)
print(job_description.preferred_skills)
print(job_description.min_experience)
print(job_description.education_requirement)
print(job_description.responsibilities)

# resume schema

class MatchResult(BaseModel):
    score:float
    details: dict
class Experience(BaseModel):
    company:str|None=None
    role:str|None=None
    duration:str|None=None
    description:str|None=None
    skills:list[str]=[]
class Resume(BaseModel):
    name:str|None=None
    email:str|None=None
    phone:str|None=None
    education:str|None=None
    experience:list[Experience]|None=None
    skills:list[str]|None=None
    match_result:MatchResult|None = None
    projects:list[dict]|None = None
    certifications:list[str]|None = None

# part 3 resume parsing
resume_schema = Resume.model_json_schema()

def parse_resume(resume_text):
    system_prompt = f"""please extract the resume information and give a json output matching this schema: {resume_schema}. Only include fields you can actually find in the text; omit fields that aren't present. important: do not return the actual schema itself do not reuturn the field like "properties","title".\n\ntext: {resume_text}"""
    user_prompt = f"""analyze the resume and extract the required information. Return the output in JSON format matching the schema provided. Only include fields that are present in the text; omit any fields that aren't found. Do not return the actual schema itself or any field names like "properties" or "title".\n\ntext: {resume_text}"""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = client.chat.completions.create(model=model, messages=messages, response_format={"type": "json_object"})
    data = json.loads(response.choices[0].message.content)
    return Resume(**data)

def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt=f"""You are a resume parser and job description analyzer. Your task is to compare the provided job description and resume, and calculate a skill match score based on the overlap of required and preferred skills. The score should be a float between 0 and 1, where 1 indicates a perfect match and 0 indicates no match. Additionally, provide details on which skills matched and which did not.

Return a json output matching this schema: {match_schema}. Only include the fields defined in the schema.

job description: {job.model_dump_json()}

resume: {resume.model_dump_json()}"""
    messages = [{"role": "system", "content": prompt}]
    response = client.chat.completions.create(model=model, messages=messages, response_format={"type": "json_object"})
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)



from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def read_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    return text

def read_resume(file_path):
    if file_path.suffix == '.pdf':
        return read_pdf(file_path)
    elif file_path.suffix == '.docx':
        return read_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")

resume_folder = Path("resumes")
all_resumes=[]
all_results=[]
for file_path in resume_folder.iterdir():
    if file_path.is_file() and file_path.suffix in ['.pdf', '.docx']:
        resume_text = read_resume(file_path)
        parsed_resume = parse_resume(resume_text)
        time.sleep(1)  # Add a delay of 1 second between requests
        result = final_score(job_description, parsed_resume)
        time.sleep(1)  # Add a delay of 1 second between requests
        print("SCORE:", result.score)
        all_results.append({
            "resume_file": file_path.name,
            "score": result.score,
            "details": result.details
        })
all_results.sort(key=lambda x: x["score"], reverse=True)
top_1=all_results[0]
worst_1=all_results[-1]

print("Top 1 Resume:", top_1)
for candidate in all_results:
    print(f"Resume: {candidate['resume_file']}, Score: {candidate['score']}")